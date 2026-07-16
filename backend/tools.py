"""File-creation tools the model can invoke via Ollama tool calling.

Every tool writes a NEW file into the user's Downloads folder: paths in
tool arguments are reduced to a bare filename (no directories, no
traversal) and existing files are never overwritten — a numeric suffix
is added instead.
"""

import re
from pathlib import Path

DOWNLOADS = Path.home() / "Downloads"

MAX_ROWS = 10_000
MAX_CONTENT_CHARS = 2_000_000

TOOL_DEFS = [
    {
        "type": "function",
        "function": {
            "name": "create_file",
            "description": (
                "Create a plain-text file (e.g. .txt, .md, .csv, .json,"
                " source code) in the user's Downloads folder."
            ),
            "parameters": {
                "type": "object",
                "required": ["filename", "content"],
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "File name including extension,"
                                       " e.g. notes.md",
                    },
                    "content": {
                        "type": "string",
                        "description": "The full text content of the file",
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": (
                "Create a Word document (.docx) in the user's Downloads"
                " folder. Paragraphs are separated by blank lines; lines"
                " starting with #, ## or ### become headings; lines"
                " starting with - become bullet points."
            ),
            "parameters": {
                "type": "object",
                "required": ["filename", "content"],
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "File name, e.g. report.docx",
                    },
                    "title": {
                        "type": "string",
                        "description": "Optional document title",
                    },
                    "content": {
                        "type": "string",
                        "description": "Document body text",
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_xlsx",
            "description": (
                "Create an Excel spreadsheet (.xlsx) in the user's"
                " Downloads folder from rows of cells. The first row is"
                " formatted as a header."
            ),
            "parameters": {
                "type": "object",
                "required": ["filename", "rows"],
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "File name, e.g. budget.xlsx",
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "Optional worksheet name",
                    },
                    "rows": {
                        "type": "array",
                        "description": "Rows of the sheet; each row is an"
                                       " array of cell values (strings or"
                                       " numbers). First row = headers.",
                        "items": {"type": "array"},
                    },
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_pdf",
            "description": (
                "Create a PDF document in the user's Downloads folder."
                " Paragraphs are separated by blank lines; lines starting"
                " with #, ## or ### become headings."
            ),
            "parameters": {
                "type": "object",
                "required": ["filename", "content"],
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "File name, e.g. summary.pdf",
                    },
                    "title": {
                        "type": "string",
                        "description": "Optional document title",
                    },
                    "content": {
                        "type": "string",
                        "description": "Document body text",
                    },
                },
            },
        },
    },
]


def _safe_new_path(filename: str, default_ext: str) -> Path:
    """Bare, sanitized filename inside Downloads that doesn't exist yet."""
    name = Path(str(filename)).name  # drop any directory components
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip(" .") or "file"
    if not name.lower().endswith(default_ext):
        name += default_ext
    DOWNLOADS.mkdir(exist_ok=True)
    path = DOWNLOADS / name
    stem, suffix = path.stem, path.suffix
    n = 1
    while path.exists():
        path = DOWNLOADS / f"{stem} ({n}){suffix}"
        n += 1
    return path


def _blocks(content: str) -> list[str]:
    return [b.strip() for b in str(content).split("\n\n") if b.strip()]


def _create_file(args: dict) -> Path:
    content = str(args.get("content", ""))
    if len(content) > MAX_CONTENT_CHARS:
        raise ValueError("Content too large")
    ext = Path(str(args.get("filename", ""))).suffix or ".txt"
    path = _safe_new_path(args.get("filename", "file"), ext)
    path.write_text(content, encoding="utf-8")
    return path


def _create_docx(args: dict) -> Path:
    from docx import Document

    doc = Document()
    if args.get("title"):
        doc.add_heading(str(args["title"]), 0)
    for block in _blocks(args.get("content", "")):
        for line in block.split("\n"):
            line = line.strip()
            if not line:
                continue
            m = re.match(r"^(#{1,3})\s+(.*)", line)
            if m:
                doc.add_heading(m.group(2), len(m.group(1)))
            elif line.startswith(("- ", "* ")):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
    path = _safe_new_path(args.get("filename", "document"), ".docx")
    doc.save(path)
    return path


def _create_xlsx(args: dict) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    rows = args.get("rows")
    if not isinstance(rows, list) or not rows:
        raise ValueError("rows must be a non-empty array of arrays")
    if len(rows) > MAX_ROWS:
        raise ValueError("Too many rows")
    wb = Workbook()
    ws = wb.active
    if args.get("sheet_name"):
        ws.title = re.sub(r"[\[\]:*?/\\]", "_", str(args["sheet_name"]))[:31]
    for row in rows:
        if not isinstance(row, list):
            row = [row]
        ws.append([
            c if isinstance(c, (int, float, bool)) or c is None else str(c)
            for c in row
        ])
    for cell in ws[1]:
        cell.font = Font(bold=True)
    path = _safe_new_path(args.get("filename", "sheet"), ".xlsx")
    wb.save(path)
    return path


def _create_pdf(args: dict) -> Path:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    # Core PDF fonts are latin-1 only; degrade unsupported characters.
    def latin(s):
        return str(s).encode("latin-1", "replace").decode("latin-1")

    if args.get("title"):
        pdf.set_font("Helvetica", "B", 18)
        pdf.multi_cell(0, 9, latin(args["title"]))
        pdf.ln(3)
    for block in _blocks(args.get("content", "")):
        m = re.match(r"^(#{1,3})\s+(.*)", block.split("\n")[0])
        if m and "\n" not in block:
            pdf.set_font("Helvetica", "B", 16 - 2 * (len(m.group(1)) - 1))
            pdf.multi_cell(0, 8, latin(m.group(2)))
            pdf.ln(1)
        else:
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 6, latin(block))
            pdf.ln(2)
    path = _safe_new_path(args.get("filename", "document"), ".pdf")
    pdf.output(str(path))
    return path


_EXECUTORS = {
    "create_file": _create_file,
    "create_docx": _create_docx,
    "create_xlsx": _create_xlsx,
    "create_pdf": _create_pdf,
}


def execute_tool(name: str, args: dict) -> dict:
    """Run a tool call; always returns a JSON-safe result dict."""
    executor = _EXECUTORS.get(name)
    if not executor:
        return {"ok": False, "error": f"Unknown tool: {name}"}
    if not isinstance(args, dict):
        return {"ok": False, "error": "Invalid tool arguments"}
    try:
        path = executor(args)
        return {"ok": True, "filename": path.name,
                "path": str(path), "folder": "Downloads"}
    except Exception as exc:  # report failure back to the model
        return {"ok": False, "error": str(exc)[:300]}
